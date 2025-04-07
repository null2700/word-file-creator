import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Inches

doc = Document()
doc.add_heading('Question Paper', 0)

questions = []

def add_question():
    q_type = question_type.get()
    content = question_text.get("1.0", tk.END).strip()

    if not content:
        messagebox.showerror("Empty", "Enter a question")
        return

    if q_type == "Fill in the Blanks":
        doc.add_paragraph(content.replace("blank", "__________"))
    elif q_type == "Table":
        try:
            rows = int(rows_entry.get())
            cols = int(cols_entry.get())
            table = doc.add_table(rows=rows, cols=cols)
            for r in range(rows):
                for c in range(cols):
                    text = table_entries[r][c].get()
                    table.cell(r, c).text = text
        except:
            messagebox.showerror("Error", "Invalid table input")
            return
    else:
        doc.add_paragraph(content)

    question_text.delete("1.0", tk.END)
    update_ui()

def update_ui():
    for widget in table_frame.winfo_children():
        widget.destroy()

    if question_type.get() == "Table":
        try:
            rows = int(rows_entry.get())
            cols = int(cols_entry.get())
            global table_entries
            table_entries = []
            for r in range(rows):
                row = []
                for c in range(cols):
                    entry = tk.Entry(table_frame, width=10)
                    entry.grid(row=r, column=c, padx=2, pady=2)
                    row.append(entry)
                table_entries.append(row)
        except:
            pass

def save_doc():
    filename = file_name.get()
    if not filename.endswith(".docx"):
        filename += ".docx"
    doc.save(filename)
    messagebox.showinfo("Saved", f"Saved as {filename}")

# GUI Setup
root = tk.Tk()
root.title("Question Paper Generator")

tk.Label(root, text="Header").grid(row=0, column=0, sticky="w")
header_entry = tk.Entry(root, width=50)
header_entry.grid(row=0, column=1, columnspan=2)
doc.add_heading(header_entry.get(), level=1)

tk.Label(root, text="Question Text").grid(row=1, column=0, sticky="nw")
question_text = tk.Text(root, height=5, width=50)
question_text.grid(row=1, column=1, columnspan=2)

tk.Label(root, text="Question Type").grid(row=2, column=0, sticky="w")
question_type = ttk.Combobox(root, values=["Simple", "Fill in the Blanks", "Table"])
question_type.current(0)
question_type.grid(row=2, column=1, sticky="w")
question_type.bind("<<ComboboxSelected>>", lambda e: update_ui())

tk.Label(root, text="Rows").grid(row=3, column=0, sticky="w")
rows_entry = tk.Entry(root, width=5)
rows_entry.insert(0, "2")
rows_entry.grid(row=3, column=1, sticky="w")

tk.Label(root, text="Columns").grid(row=3, column=1, sticky="e")
cols_entry = tk.Entry(root, width=5)
cols_entry.insert(0, "2")
cols_entry.grid(row=3, column=2, sticky="w")

table_frame = tk.Frame(root)
table_frame.grid(row=4, column=0, columnspan=3)

tk.Button(root, text="Add Question", command=add_question).grid(row=5, column=0, pady=10)
tk.Label(root, text="Save As").grid(row=6, column=0, sticky="e")
file_name = tk.Entry(root, width=30)
file_name.insert(0, "Question_Paper.docx")
file_name.grid(row=6, column=1)
tk.Button(root, text="Save Document", command=save_doc).grid(row=6, column=2)

root.mainloop()
